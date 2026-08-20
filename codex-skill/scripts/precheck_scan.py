from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
EXTRACTOR_VERSION = "1.1.0"
SOURCE_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".zip"}
EXCLUDED_DIR_PATTERNS = (
    re.compile(r"^99_プレチェック出力(?:_|$)"),
    re.compile(r"^tmp$", re.I),
    re.compile(r"^output(?:s)?$", re.I),
    re.compile(r"^cache$", re.I),
    re.compile(r"^\.precheck-cache$", re.I),
    re.compile(r"^\.git$", re.I),
)
DATE_PATTERN = re.compile(
    r"(?P<year>20\d{2})\s*[年/.-]\s*(?P<month>1[0-2]|0?[1-9])\s*[月/.-]\s*(?P<day>3[01]|[12]\d|0?[1-9])\s*日?"
)
MONTH_PATTERN = re.compile(r"(?P<months>\d{1,3})\s*(?:か|ヶ|ケ|箇)?月")
TARGET_PATTERN = re.compile(r"(?P<count>\d{1,6})\s*(?:例|名|人)")
TIMELINE_TERMS = (
    "研究期間", "実施期間", "研究終了", "登録期間", "エントリー", "追跡", "観察期間",
    "解析", "統計解析", "症例数", "対象者数", "予定対象者数",
)
EMAIL_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_PATTERN = re.compile(r"(?<!\d)0\d{1,4}[-‐‑–—ー ]?\d{1,4}[-‐‑–—ー ]?\d{3,4}(?!\d)")
LABELED_FACT_TERMS = {
    "study_titles": ("研究課題名", "課題名"),
    "people_and_roles": ("研究責任者", "研究分担者", "研究協力者", "個人情報管理責任者", "連絡担当者", "氏名"),
    "affiliations_and_sites": ("所属", "実施施設", "研究機関", "病院", "大学", "センター"),
    "storage_periods": ("保管期間", "保存期間", "研究終了後", "結果公表後"),
}
RULE_FILE_GROUPS = (
    ("R1_項目間整合ルール.md", "r1-cross-document-consistency.md"),
    ("R2_必須記載チェックリスト.md", "r2-required-content.md"),
)


def default_rules_dir() -> Path:
    root = Path(__file__).resolve().parents[1]
    packaged = root / "references" / "rules"
    return packaged if packaged.is_dir() else root / "rules"


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def is_excluded(relative: Path) -> bool:
    return any(any(pattern.search(part) for pattern in EXCLUDED_DIR_PATTERNS) for part in relative.parts[:-1])


def discover_sources(study: Path) -> tuple[list[Path], list[str]]:
    sources: list[Path] = []
    warnings: list[str] = []
    for path in sorted(study.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SOURCE_SUFFIXES:
            continue
        rel = path.relative_to(study)
        if is_excluded(rel):
            if path.suffix.lower() in {".pdf", ".docx"}:
                warnings.append(f"excluded-source-like-file: {rel}")
            continue
        sources.append(path)
    return sources, warnings


def extract_pdf(path: Path) -> dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for number, page in enumerate(reader.pages, 1):
        pages.append({"page": number, "text": page.extract_text() or ""})
    return {"kind": "pdf", "pages": pages}


def extract_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    blocks: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        if paragraph.text.strip():
            blocks.append(f"[P{index}] {paragraph.text}")
    for table_no, table in enumerate(document.tables, 1):
        for row_no, row in enumerate(table.rows, 1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            blocks.append(f"[T{table_no}R{row_no}] " + " | ".join(cells))
    return {"kind": "docx", "pages": [{"page": None, "text": "\n".join(blocks)}]}


def extract_plain(path: Path) -> dict[str, Any]:
    return {"kind": path.suffix.lower().lstrip("."), "pages": [{"page": None, "text": path.read_text(encoding="utf-8", errors="replace")}]}


def extract_source(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".md", ".csv"}:
        return extract_plain(path)
    return {"kind": suffix.lstrip("."), "pages": [], "note": "binary inventory only"}


def cache_key(file_hash: str) -> str:
    return hashlib.sha256(f"{EXTRACTOR_VERSION}:{file_hash}".encode()).hexdigest()


def get_extraction(path: Path, file_hash: str, cache_dir: Path, force: bool) -> tuple[dict[str, Any], bool]:
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / f"{cache_key(file_hash)}.json"
    if cache_file.exists() and not force:
        return json.loads(cache_file.read_text(encoding="utf-8")), True
    payload = extract_source(path)
    payload.update({"extractor_version": EXTRACTOR_VERSION, "sha256": file_hash})
    cache_file.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload, False


def compact_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def timeline_evidence(extractions: list[dict[str, Any]]) -> dict[str, Any]:
    evidence: list[dict[str, Any]] = []
    for item in extractions:
        for page in item["extraction"].get("pages", []):
            lines = (page.get("text") or "").splitlines()
            for line_no, raw in enumerate(lines, 1):
                line = compact_line(raw)
                if not line or not any(term in line for term in TIMELINE_TERMS):
                    continue
                start = max(0, line_no - 2)
                end = min(len(lines), line_no + 2)
                context = compact_line(" ".join(lines[start:end]))
                dates = [m.group(0) for m in DATE_PATTERN.finditer(context)]
                months = [int(m.group("months")) for m in MONTH_PATTERN.finditer(context)]
                targets = [int(m.group("count")) for m in TARGET_PATTERN.finditer(context)]
                evidence.append({
                    "source": item["path"], "page": page.get("page"), "line": line_no,
                    "text": line[:500], "context": context[:1000], "dates": dates, "months": months, "targets": targets,
                })
    return {
        "schema_version": 1,
        "status": "requires_reviewer_confirmation",
        "mandatory_questions": [
            "研究開始日と研究終了日は全資料で一致するか",
            "明記された研究期間（月数）は年月日の差と一致するか",
            "最終登録後の追跡期間を研究終了日までに確保できるか",
            "追跡終了後から統計解析終了までの期間を確保できるか",
            "目標症例数と登録見込みから登録完了時期は現実的か",
        ],
        "evidence": evidence,
        "review_fields": {
            "research_start": None, "research_end": None, "stated_total_months": None,
            "target_count": None, "expected_enrollment_end": None, "followup_months": None,
            "analysis_months": None, "latest_followup_end": None, "latest_analysis_end": None,
            "analysis_margin_days": None, "finding": None, "assumptions": [],
        },
    }


def compact_facts(extractions: list[dict[str, Any]]) -> dict[str, Any]:
    buckets: dict[str, dict[str, dict[str, Any]]] = {name: {} for name in LABELED_FACT_TERMS}
    buckets.update({"emails": {}, "phone_numbers": {}, "target_counts": {}, "dates": {}})
    detected_items: set[int] = set()
    item_evidence: list[dict[str, Any]] = []

    def record(category: str, value: str, source: str, page: int | None, context: str) -> None:
        normalized = normalize_for_scan(value)
        if category == "emails":
            normalized = value.lower()
        elif category == "phone_numbers":
            normalized = re.sub(r"\D", "", value)
        if not normalized:
            return
        row = buckets[category].setdefault(normalized, {
            "value": compact_line(value)[:500], "normalized": normalized[:500],
            "source": source, "page": page, "context": compact_line(context)[:700], "occurrences": 0,
        })
        row["occurrences"] += 1

    for item in extractions:
        for page in item["extraction"].get("pages", []):
            for raw in (page.get("text") or "").splitlines():
                line = compact_line(raw)
                if not line:
                    continue
                for category, terms in LABELED_FACT_TERMS.items():
                    if any(term in line for term in terms):
                        record(category, line, item["path"], page.get("page"), line)
                for match in EMAIL_PATTERN.finditer(line):
                    record("emails", match.group(0), item["path"], page.get("page"), line)
                for match in PHONE_PATTERN.finditer(line):
                    record("phone_numbers", match.group(0), item["path"], page.get("page"), line)
                for match in TARGET_PATTERN.finditer(line):
                    record("target_counts", match.group(0), item["path"], page.get("page"), line)
                for match in DATE_PATTERN.finditer(line):
                    record("dates", match.group(0), item["path"], page.get("page"), line)
                numbers = {
                    int(value) for value in re.findall(r"(?:^|[^0-9])(?:項目\s*)?([1-9]|[1-4]\d|5[01])\s*[.．:：]", raw)
                }
                for number in sorted(numbers - detected_items):
                    detected_items.add(number)
                    item_evidence.append({"item": number, "source": item["path"], "page": page.get("page"), "text": line[:300]})
    return {
        "schema_version": 1,
        "fact_candidates": {name: list(values.values()) for name, values in buckets.items()},
        "application_items": {
            "detected": sorted(detected_items),
            "missing_candidates": sorted(set(range(1, 52)) - detected_items),
            "evidence": item_evidence,
            "review_note": "番号検出は候補。申請システム本文で1〜51の欠番を確認する。",
        },
    }


def rule_ids(rules_dir: Path) -> list[str]:
    found: list[str] = []
    pattern = re.compile(r"^\|\s*((?:G|C)\d+-[0-9A-Za-z]+)\s*\|")
    for alternatives in RULE_FILE_GROUPS:
        path = next((rules_dir / name for name in alternatives if (rules_dir / name).is_file()), None)
        if path is None:
            raise FileNotFoundError(f"Missing rule file under {rules_dir}: one of {alternatives}")
        for line in path.read_text(encoding="utf-8").splitlines():
            match = pattern.match(line)
            if match and match.group(1) not in found:
                found.append(match.group(1))
    return found


def r3_hits(rules_dir: Path, extractions: list[dict[str, Any]]) -> dict[str, Any]:
    alternatives = ("R3_書式残骸・表記ルール.md", "r3-format-and-word-qa.md")
    rule_file = next((rules_dir / name for name in alternatives if (rules_dir / name).is_file()), None)
    if rule_file is None:
        raise FileNotFoundError(f"Missing R3 rule file under {rules_dir}: one of {alternatives}")
    machine_patterns: list[tuple[str, str]] = []
    unscanned: list[str] = []
    for line in rule_file.read_text(encoding="utf-8").splitlines():
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if len(cells) < 2 or not re.fullmatch(r"T[12]-\d+", cells[0]):
            continue
        rid, specification = cells[0], cells[1]
        tokens = re.findall(r"`([^`]+)`", specification)
        if rid == "T1-4":
            tokens = [r"(?:\[|［)?試料(?:および情報)?・(?:試料・)?情報(?:\]|］)?"]
        if rid.startswith("T2-") and not tokens:
            tokens = [part.strip() for part in specification.split("／") if part.strip()]
            if rid == "T2-1":
                tokens = [token.replace("（法人名として）", "") for token in tokens]
        if not tokens and rid != "T1-13":
            unscanned.append(rid)
        for token in tokens:
            machine_patterns.append((rid, token))
    hits: list[dict[str, Any]] = []
    invalid_patterns: list[dict[str, str]] = []
    for rid, token in machine_patterns:
        try:
            pattern = re.compile(token)
        except re.error as exc:
            invalid_patterns.append({"rule_id": rid, "pattern": token, "error": str(exc)})
            continue
        for item in extractions:
            for page in item["extraction"].get("pages", []):
                raw = page.get("text") or ""
                search_versions = (("normalized", normalize_for_scan(raw)), ("raw", raw))
                seen: set[str] = set()
                for mode, searchable in search_versions:
                    for match in pattern.finditer(searchable):
                        key = match.group(0)
                        if key in seen:
                            continue
                        seen.add(key)
                        hits.append({
                            "rule_id": rid, "source": item["path"], "page": page.get("page"),
                            "pattern": token, "mode": mode, "matched": key[:200],
                        })
    # T1-13 is structural rather than textual. DOCX extraction preserves table/row
    # markers, so blank rows can be checked only in tables whose header contains
    # both affiliation and name. PDF blank rows remain part of visual QA.
    for item in extractions:
        if item["extraction"].get("kind") != "docx":
            continue
        text = "\n".join(page.get("text") or "" for page in item["extraction"].get("pages", []))
        lines = text.splitlines()
        eligible_tables: set[str] = set()
        for line in lines:
            marker = re.match(r"\[(T\d+)R\d+\]\s*(.*)", line)
            if marker and "所属" in marker.group(2) and "氏名" in marker.group(2):
                eligible_tables.add(marker.group(1))
        for line in lines:
            marker = re.match(r"\[(T\d+)R(\d+)\]\s*(.*)", line)
            if not marker or marker.group(1) not in eligible_tables:
                continue
            cells = [cell.strip() for cell in marker.group(3).split("|")]
            if len(cells) >= 2 and all(not cell for cell in cells):
                hits.append({
                    "rule_id": "T1-13", "source": item["path"], "page": None,
                    "pattern": "blank affiliation/name table row", "mode": "docx-structure",
                    "matched": f"{marker.group(1)}R{marker.group(2)}",
                })
    return {
        "schema_version": 1, "rule_file": str(rule_file), "hits": hits,
        "unscanned_rules": sorted(set(unscanned)), "invalid_patterns": invalid_patterns,
    }


def normalize_for_scan(text: str) -> str:
    return re.sub(r"[\s　]+", "", text or "")


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract one study once and create a compact precheck fact pack.")
    parser.add_argument("study", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--cache", type=Path, default=Path(".precheck-cache"))
    parser.add_argument("--rules", type=Path, default=default_rules_dir())
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args()
    study = args.study.resolve()
    out = args.out.resolve()
    out.mkdir(parents=True, exist_ok=True)
    sources, warnings = discover_sources(study)
    if not sources:
        raise SystemExit(f"No source files found under {study}")

    manifest_files: list[dict[str, Any]] = []
    extracted: list[dict[str, Any]] = []
    for path in sources:
        stat = path.stat()
        file_hash = sha256_file(path)
        extraction, hit = get_extraction(path, file_hash, args.cache.resolve(), args.force)
        rel = str(path.relative_to(study))
        manifest_files.append({
            "path": rel, "suffix": path.suffix.lower(), "size": stat.st_size,
            "mtime_ns": stat.st_mtime_ns, "sha256": file_hash, "cache_hit": hit,
        })
        extracted.append({"path": rel, "extraction": extraction})

    manifest = {
        "schema_version": 1, "created_at": utc_now(), "study": str(study),
        "extractor_version": EXTRACTOR_VERSION, "excluded_directory_policy": [p.pattern for p in EXCLUDED_DIR_PATTERNS],
        "warnings": warnings, "files": manifest_files,
    }
    (out / "00_source_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    (out / "00_extracted_pages.json").write_text(json.dumps(extracted, ensure_ascii=False, indent=2), encoding="utf-8")
    timeline = timeline_evidence(extracted)
    (out / "00_timeline_facts.json").write_text(json.dumps(timeline, ensure_ascii=False, indent=2), encoding="utf-8")
    residue = r3_hits(args.rules.resolve(), extracted)
    (out / "00_r3_hits.json").write_text(json.dumps(residue, ensure_ascii=False, indent=2), encoding="utf-8")
    ledger = {
        "schema_version": 2, "study": study.name,
        "allowed_statuses": ["na", "satisfied", "candidate", "human"],
        "allowed_evidence_statuses": ["confirmed", "inferred", "human"],
        "allowed_return_routes": [
            "applicant_return", "attachment_return", "researcher_profile",
            "secretariat_internal", "human_decision", "resolved_elsewhere",
        ],
        "rules": [
            {
                "rule_id": rid, "status": None, "evidence": [], "note": "",
                "root_issue_id": None, "linked_items": [],
            }
            for rid in rule_ids(args.rules.resolve())
        ],
        "dispositions": [],
        "disposition_schema": {
            "finding_id": "案件内で一意のID",
            "finding_type": "A|B",
            "source_rule_ids": ["G1-1"],
            "item_numbers": [1],
            "document": "B欄候補の資料名。A欄はnull",
            "evidence_status": "confirmed|inferred|human",
            "return_route": "applicant_return|attachment_return|researcher_profile|secretariat_internal|human_decision|resolved_elsewhere",
            "action_owner": "申請者・各研究者本人・事務局・最終確認者など",
            "included_in_return": True,
            "final_text": "04とWordへ反映する確定文。非返却時は空文字",
            "disposition_reason": "採用・集約・別経路対応・保留等の理由",
            "root_issue_id": "同じ原因から派生する指摘を束ねるID",
        },
    }
    (out / "00_rule_ledger.json").write_text(json.dumps(ledger, ensure_ascii=False, indent=2), encoding="utf-8")
    facts = compact_facts(extracted)
    summary = {
        "schema_version": 2,
        "study": study.name, "source_count": len(sources), "cache_hits": sum(bool(f["cache_hit"]) for f in manifest_files),
        "timeline_evidence_count": len(timeline["evidence"]), "r3_hit_count": len(residue["hits"]),
        "rule_count": len(ledger["rules"]), "warnings": warnings,
        "fact_candidates": facts["fact_candidates"],
        "application_items": facts["application_items"],
        "r3_hits": residue["hits"],
        "review_note": "抽出候補であり、原資料頁による照合と文書間比較を省略しない。",
    }
    (out / "00_fact_pack.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    console_summary = {
        "study": study.name,
        "source_count": len(sources),
        "cache_hits": sum(bool(f["cache_hit"]) for f in manifest_files),
        "timeline_evidence_count": len(timeline["evidence"]),
        "r3_hit_count": len(residue["hits"]),
        "rule_count": len(ledger["rules"]),
        "warning_count": len(warnings),
        "output": str(out),
    }
    print(json.dumps(console_summary, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
