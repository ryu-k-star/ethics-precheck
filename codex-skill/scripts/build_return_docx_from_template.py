from __future__ import annotations

import argparse
import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt


RETURN_PREFIX = "★倫理申請_修正対応依頼_"
OUTPUT_FOLDER_PREFIX = "99_プレチェック出力"
A_TABLE_HEADER_PREFIX = "| 項目 | 対応要否 |"
BOTTOM_SECTION_MARKER = "A．倫理審査申請システム質疑事項"
EXPECTED_MAIN_TABLE_ROWS = 52
WINDOWS_FILENAME_TRANSLATION = str.maketrans(
    {
        '"': "”",
        "*": "＊",
        "/": "／",
        ":": "：",
        "<": "＜",
        ">": "＞",
        "?": "？",
        "\\": "￥",
        "|": "｜",
    }
)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def find_a_table_end(text: str) -> int:
    offset = 0
    in_table = False
    saw_row = False
    for raw_line in text.splitlines(keepends=True):
        line = raw_line.rstrip("\r\n")
        if not in_table:
            if line.startswith(A_TABLE_HEADER_PREFIX):
                in_table = True
            offset += len(raw_line)
            continue

        if line.startswith("|---"):
            offset += len(raw_line)
            continue
        if line.startswith("|"):
            saw_row = True
            offset += len(raw_line)
            continue
        if saw_row:
            return offset
        offset += len(raw_line)

    if not in_table or not saw_row:
        raise ValueError("No complete A-section markdown table found in draft.")
    return len(text)


def split_bottom_sections(text: str) -> tuple[str, str]:
    table_end = find_a_table_end(text)
    matches = [
        match.start()
        for match in re.finditer(
            re.escape(BOTTOM_SECTION_MARKER), text[table_end:]
        )
    ]
    if len(matches) != 1:
        raise ValueError(
            "Draft must contain exactly one footer marker after the A-section table: "
            f"{BOTTOM_SECTION_MARKER}; found {len(matches)}."
        )
    idx = table_end + matches[0]
    return text[:idx].strip(), text[idx:].strip()


def extract_title(text: str) -> str:
    for line in text.splitlines():
        if line.startswith("研究課題名"):
            return line.strip()
    return "研究課題名【】"


def extract_study_title(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped.startswith("研究課題名"):
            continue
        value = stripped[len("研究課題名") :].strip()
        if value.startswith("【") and value.endswith("】"):
            value = value[1:-1].strip()
        else:
            value = value.lstrip("：:").strip()
        if value:
            return value
    raise ValueError("Draft must contain a non-empty Japanese study title.")


def sanitize_filename_component(value: str) -> str:
    sanitized = value.translate(WINDOWS_FILENAME_TRANSLATION)
    sanitized = "".join(" " if ord(char) < 32 else char for char in sanitized)
    sanitized = re.sub(r"\s+", " ", sanitized).strip(" .")
    if not sanitized:
        raise ValueError("Study title becomes empty after filename sanitization.")
    return sanitized


def validate_output_path(output: Path, study_title: str) -> None:
    output_folder = output.parent.name
    if output_folder != OUTPUT_FOLDER_PREFIX and not output_folder.startswith(
        f"{OUTPUT_FOLDER_PREFIX}_"
    ):
        raise ValueError(
            f"Return DOCX must be saved in {OUTPUT_FOLDER_PREFIX}/ or "
            f"{OUTPUT_FOLDER_PREFIX}_<label>/; got: {output_folder}"
        )

    safe_title = sanitize_filename_component(study_title)
    expected_pattern = re.compile(
        rf"^{re.escape(RETURN_PREFIX + safe_title)}_(\d{{4}})\.docx$",
        re.IGNORECASE,
    )
    if not expected_pattern.fullmatch(output.name):
        expected = f"{RETURN_PREFIX}{safe_title}_作成日4桁.docx"
        raise ValueError(
            "Return DOCX filename does not follow the required Japanese naming rule. "
            f"Expected: {expected}; got: {output.name}"
        )


def parse_markdown_table(text: str) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    in_table = False
    for line in text.splitlines():
        if line.startswith(A_TABLE_HEADER_PREFIX):
            in_table = True
            continue
        if not in_table:
            continue
        if line.startswith("|---"):
            continue
        if not line.startswith("|"):
            if rows:
                break
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if len(cells) < 3:
            continue
        rows.append((cells[0], cells[1], cells[2].replace("<br>", "\n")))
    return rows


def normalize_item_label(value: str) -> str:
    return re.sub(r"\s+", "", value)


def validate_template_and_rows(doc, rows: list[tuple[str, str, str]]) -> None:
    if len(doc.tables) != 2:
        raise ValueError(
            f"Template must contain exactly two tables; got {len(doc.tables)}."
        )
    if len(doc.tables[0].rows) != EXPECTED_MAIN_TABLE_ROWS:
        raise ValueError(
            "Template main table must contain exactly "
            f"{EXPECTED_MAIN_TABLE_ROWS} rows; got {len(doc.tables[0].rows)}."
        )
    if any(len(row.cells) != 2 for row in doc.tables[0].rows):
        raise ValueError("Every row in the template main table must have two cells.")
    if len(doc.tables[1].rows) != 1 or len(doc.tables[1].rows[0].cells) != 1:
        raise ValueError("Template footer table must be exactly one row by one cell.")
    if len(rows) != EXPECTED_MAIN_TABLE_ROWS:
        raise ValueError(
            "Draft A-section must contain exactly "
            f"{EXPECTED_MAIN_TABLE_ROWS} rows; got {len(rows)}."
        )

    expected_labels = [
        normalize_item_label(row.cells[0].text) for row in doc.tables[0].rows
    ]
    actual_labels = [normalize_item_label(item) for item, _status, _comment in rows]
    if actual_labels != expected_labels:
        mismatch = next(
            index
            for index, (actual, expected) in enumerate(
                zip(actual_labels, expected_labels), start=1
            )
            if actual != expected
        )
        raise ValueError(
            "Draft A-section row labels or order do not match the template at "
            f"row {mismatch}: expected {expected_labels[mismatch - 1]!r}, "
            f"got {actual_labels[mismatch - 1]!r}."
        )

    allowed_statuses = {"要対応", "対応不要"}
    invalid_statuses = sorted(
        {status for _item, status, _comment in rows if status not in allowed_statuses}
    )
    if invalid_statuses:
        raise ValueError(
            "Draft A-section contains unsupported status values: "
            + ", ".join(invalid_statuses)
        )


def clear_cell(cell) -> None:
    cell.text = ""


def write_footer(cell, text: str) -> None:
    cell.text = ""
    for index, line in enumerate(text.splitlines()):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        run = paragraph.add_run(line)
        if line.startswith(("A．", "B．", "C．", "【", "★")):
            run.bold = True
        if line.startswith(("A．", "B．", "C．", "【")):
            paragraph.paragraph_format.keep_with_next = True


def build_docx(template: Path, draft: Path, output: Path) -> None:
    draft_text = read_text(draft)
    study_title = extract_study_title(draft_text)
    validate_output_path(output, study_title)

    rows = parse_markdown_table(draft_text)
    if not rows:
        raise ValueError("No A-section markdown table found in draft.")
    source_doc = Document(template)
    validate_template_and_rows(source_doc, rows)
    top_text, bottom_text = split_bottom_sections(draft_text)
    title = extract_title(top_text)

    output.parent.mkdir(parents=True, exist_ok=True)
    temporary_output = output.with_name(f".{output.stem}.tmp.docx")
    try:
        shutil.copy2(template, temporary_output)
        doc = Document(temporary_output)
        doc.styles["Normal"].font.name = "Yu Gothic"
        doc.styles["Normal"].font.size = Pt(10.5)

        if len(doc.paragraphs) > 1:
            doc.paragraphs[1].text = title

        table = doc.tables[0]
        for row_idx in range(len(table.rows)):
            clear_cell(table.cell(row_idx, 1))

        for idx, (_item, status, comment) in enumerate(rows):
            if status == "要対応" and comment:
                table.cell(idx, 1).text = comment

        write_footer(doc.tables[1].cell(0, 0), bottom_text)
        doc.save(temporary_output)
        os.replace(temporary_output, output)
    finally:
        temporary_output.unlink(missing_ok=True)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Create a return DOCX while preserving the official template tables."
    )
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--draft", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build_docx(args.template, args.draft, args.output)


if __name__ == "__main__":
    main()
