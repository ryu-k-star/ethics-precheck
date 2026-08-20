from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document


RULE_ID = re.compile(r"^\|\s*((?:G|C)\d+-[0-9A-Za-z]+)\s*\|")
A_ROW = re.compile(r"^\|\s*([^|]+?)\s*\|\s*(要対応|対応不要|人間確認|要検討)\s*\|\s*(.*?)\s*\|$")
RULE_FILE_GROUPS = (
    ("R1_項目間整合ルール.md", "r1-cross-document-consistency.md"),
    ("R2_必須記載チェックリスト.md", "r2-required-content.md"),
)
RETURN_PREFIX = "★倫理申請_修正対応依頼_"
WINDOWS_FILENAME_TRANSLATION = str.maketrans({
    '"': "”", "*": "＊", "/": "／", ":": "：", "<": "＜", ">": "＞",
    "?": "？", "\\": "￥", "|": "｜",
})


def default_rules_dir() -> Path:
    root = Path(__file__).resolve().parents[1]
    packaged = root / "references" / "rules"
    return packaged if packaged.is_dir() else root / "rules"


def expected_rule_ids(rules_dir: Path) -> list[str]:
    result: list[str] = []
    for alternatives in RULE_FILE_GROUPS:
        path = next((rules_dir / name for name in alternatives if (rules_dir / name).is_file()), None)
        if path is None:
            raise FileNotFoundError(f"Missing rule file under {rules_dir}: one of {alternatives}")
        for line in path.read_text(encoding="utf-8").splitlines():
            match = RULE_ID.match(line)
            if match and match.group(1) not in result:
                result.append(match.group(1))
    return result


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def item_numbers(value: str) -> set[int]:
    return {int(x) for x in re.findall(r"(?<!\d)([1-9]|[1-4]\d|5[01])(?!\d)", value)}


def docx_xml_flags(path: Path) -> dict[str, int]:
    comments = tracked = 0
    with zipfile.ZipFile(path) as package:
        names = package.namelist()
        comments = sum(1 for name in names if name.startswith("word/comments"))
        for name in names:
            if name.startswith("word/") and name.endswith(".xml"):
                data = package.read(name)
                tracked += len(re.findall(br"<w:(?:ins|del)(?:\s|>)", data))
    return {"comment_parts": comments, "tracked_change_tags": tracked}


def add(checks: list[dict[str, Any]], check_id: str, passed: bool, detail: str) -> None:
    checks.append({"id": check_id, "passed": bool(passed), "detail": detail})


def row_item_number(label: str) -> int | None:
    match = re.match(r"\s*(\d{1,2})\.", label or "")
    return int(match.group(1)) if match else None


def expected_word_name(draft: str, actual_name: str) -> str | None:
    study_title: str | None = None
    for line in draft.splitlines():
        stripped = line.strip()
        if not stripped.startswith("研究課題名"):
            continue
        value = stripped[len("研究課題名"):].strip()
        if value.startswith("【") and value.endswith("】"):
            value = value[1:-1].strip()
        else:
            value = value.lstrip("：:").strip()
        if value:
            study_title = value
            break
    date_match = re.search(r"_(\d{4})\.docx$", actual_name, re.IGNORECASE)
    if not study_title or not date_match:
        return None
    safe_title = study_title.translate(WINDOWS_FILENAME_TRANSLATION)
    safe_title = "".join(" " if ord(char) < 32 else char for char in safe_title)
    safe_title = re.sub(r"\s+", " ", safe_title).strip(" .")
    return f"{RETURN_PREFIX}{safe_title}_{date_match.group(1)}.docx" if safe_title else None


def markdown_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_candidates(review: str) -> tuple[list[dict[str, Any]], list[str]]:
    section: str | None = None
    findings: list[dict[str, Any]] = []
    errors: list[str] = []
    for line in review.splitlines():
        stripped = line.strip()
        if stripped == "## A欄候補":
            section = "A"
            continue
        if stripped == "## B欄候補":
            section = "B"
            continue
        if stripped.startswith("## "):
            section = None
            continue
        if section not in {"A", "B"} or not stripped.startswith("|"):
            continue
        cells = markdown_cells(stripped)
        if not cells or cells[0] in {"Finding ID", "指摘ID", "---"} or set(cells[0]) == {"-"}:
            continue
        if len(cells) < 5:
            errors.append(f"{section}:invalid-column-count:{stripped[:120]}")
            continue
        finding_id, rule_text, target, pages, content = cells[:5]
        if not re.fullmatch(rf"{section}-\d{{3,}}", finding_id):
            errors.append(f"invalid-finding-id:{finding_id}")
        findings.append({
            "finding_id": finding_id, "finding_type": section,
            "source_rule_ids": re.findall(r"(?:G|C)\d+-[0-9A-Za-z]+", rule_text),
            "item_numbers": sorted(item_numbers(target)) if section == "A" else [],
            "document": target if section == "B" else None,
            "pages": pages, "content": content,
        })
    ids = [row["finding_id"] for row in findings]
    errors.extend(f"duplicate-finding-id:{fid}" for fid in sorted({fid for fid in ids if ids.count(fid) > 1}))
    return findings, errors


def parse_draft(draft: str) -> tuple[dict[int, str], str, str]:
    comments: dict[int, str] = {}
    lines = draft.splitlines()
    for line in lines:
        match = A_ROW.match(line)
        if not match or match.group(2) != "要対応":
            continue
        for number in item_numbers(match.group(1)):
            comments[number] = match.group(3).replace("<br>", "\n").strip()
    marker = "A．倫理審査申請システム質疑事項"
    indices = [index for index, line in enumerate(lines) if marker in line]
    footer = "\n".join(lines[indices[-1]:]) if indices else ""
    footer_lines = footer.splitlines()
    b_indices = [index for index, line in enumerate(footer_lines) if line.strip().startswith("B．")]
    c_indices = [index for index, line in enumerate(footer_lines) if line.strip().startswith("C．")]
    b_start = b_indices[0] if b_indices else None
    b_end = next((index for index in c_indices if b_start is not None and index > b_start), len(footer_lines))
    b_section = "\n".join(footer_lines[b_start:b_end]) if b_start is not None else ""
    return comments, b_section, footer


def verify(args: argparse.Namespace) -> dict[str, Any]:
    checks: list[dict[str, Any]] = []
    expected = expected_rule_ids(args.rules)
    ledger = json.loads(args.ledger.read_text(encoding="utf-8"))
    by_id = {row.get("rule_id"): row for row in ledger.get("rules", [])}
    missing = [rid for rid in expected if rid not in by_id]
    invalid = [rid for rid, row in by_id.items() if row.get("status") not in {"na", "satisfied", "candidate", "human"}]
    add(checks, "rules.all_present", not missing, f"missing={missing}")
    add(checks, "rules.all_classified", not invalid, f"unclassified_or_invalid={invalid}")
    duration_ids = [rid for rid in expected if rid.startswith("G6-")]
    duration_bad = [rid for rid in duration_ids if rid not in by_id or by_id[rid].get("status") not in {"na", "satisfied", "candidate", "human"}]
    add(checks, "duration.g6_complete", not duration_bad, f"required={duration_ids}; bad={duration_bad}")
    non_applicable_error = [rid for rid in ("G6-1", "G6-2") if rid in by_id and by_id[rid].get("status") == "na"]
    add(checks, "duration.core_applicable", not non_applicable_error, f"must_not_be_na={non_applicable_error}")
    timeline = json.loads(args.timeline.read_text(encoding="utf-8"))
    review_fields = timeline.get("review_fields", {})
    required_timeline_fields = {
        "research_start", "research_end", "stated_total_months", "target_count",
        "expected_enrollment_end", "followup_months", "analysis_months",
        "latest_followup_end", "latest_analysis_end", "analysis_margin_days",
        "finding", "assumptions",
    }
    missing_timeline_fields = sorted(required_timeline_fields - set(review_fields))
    timeline_ok = (timeline.get("status") == "reviewed" and not missing_timeline_fields
                   and bool(str(review_fields.get("finding") or "").strip())
                   and isinstance(review_fields.get("assumptions"), list))
    add(checks, "duration.timeline_reviewed", timeline_ok, f"status={timeline.get('status')}; missing_fields={missing_timeline_fields}; finding_present={bool(str(review_fields.get('finding') or '').strip())}")
    g67 = by_id.get("G6-7", {}).get("status")
    if g67 not in {"candidate", "human"}:
        missing_phase_check = [rid for rid in ("G6-5", "G6-6") if by_id.get(rid, {}).get("status") == "na"]
    else:
        missing_phase_check = []
    add(checks, "duration.phases_assessed", not missing_phase_check, f"G6-7={g67}; invalid_na={missing_phase_check}")

    review = args.review.read_text(encoding="utf-8")
    draft = args.draft.read_text(encoding="utf-8")
    factcheck = args.factcheck.read_text(encoding="utf-8")
    findings, candidate_errors = parse_candidates(review)
    unknown_candidate_rules = sorted({rid for finding in findings for rid in finding["source_rule_ids"] if rid not in expected})
    empty_candidate_rules = [finding["finding_id"] for finding in findings if not finding["source_rule_ids"]]
    candidate_errors.extend(f"unknown-rule-id:{rid}" for rid in unknown_candidate_rules)
    candidate_errors.extend(f"missing-rule-id:{fid}" for fid in empty_candidate_rules)
    add(checks, "findings.schema", not candidate_errors, f"count={len(findings)}; errors={candidate_errors}")
    draft_comments, draft_b, draft_footer = parse_draft(draft)
    allowed_evidence = set(ledger.get("allowed_evidence_statuses", ["confirmed", "inferred", "human"]))
    non_return_routes = {
        "researcher_profile", "secretariat_internal", "human_decision", "resolved_elsewhere",
    }
    dispositions = ledger.get("dispositions", [])
    allowed_routes = set(ledger.get("allowed_return_routes", {
        "applicant_return", "attachment_return", "researcher_profile",
        "secretariat_internal", "human_decision", "resolved_elsewhere",
    }))
    invalid_dispositions: list[str] = []
    dispositions_by_id: dict[str, dict[str, Any]] = {}
    for row in dispositions:
        raw_finding_id = row.get("finding_id")
        finding_id = str(raw_finding_id or "<missing>")
        reason = str(row.get("disposition_reason") or "").strip()
        evidence_status = row.get("evidence_status")
        route = row.get("return_route")
        included = row.get("included_in_return")
        numbers = {int(n) for n in row.get("item_numbers", []) if isinstance(n, int) or str(n).isdigit()}
        action_owner = str(row.get("action_owner") or "").strip()
        finding_type = row.get("finding_type")
        final_text = str(row.get("final_text") or "").strip()
        document_name = str(row.get("document") or "").strip()
        source_rule_ids = row.get("source_rule_ids")
        shape_ok = (finding_type == "A" and bool(numbers)) or (finding_type == "B" and bool(document_name))
        if (not raw_finding_id or not reason or evidence_status not in allowed_evidence or route not in allowed_routes
                or not action_owner or not isinstance(included, bool) or not shape_ok or not isinstance(source_rule_ids, list) or not source_rule_ids
                or any(rid not in expected for rid in source_rule_ids)
                or (included and not final_text)):
            invalid_dispositions.append(finding_id)
            continue
        if finding_id in dispositions_by_id:
            invalid_dispositions.append(f"duplicate:{finding_id}")
        dispositions_by_id[finding_id] = row
    add(checks, "disposition.valid", not invalid_dispositions, f"invalid={invalid_dispositions}")
    expected_ids = {row["finding_id"] for row in findings}
    disposition_ids = set(dispositions_by_id)
    coverage_errors = sorted(expected_ids - disposition_ids) + [f"extra:{fid}" for fid in sorted(disposition_ids - expected_ids)]
    add(checks, "disposition.coverage", not coverage_errors, f"errors={coverage_errors}")
    transfer_errors: list[str] = []
    routing_errors: list[str] = []
    for finding in findings:
        fid = finding["finding_id"]
        disposition = dispositions_by_id.get(fid)
        if not disposition:
            continue
        if disposition.get("finding_type") != finding["finding_type"]:
            transfer_errors.append(f"{fid}:type-mismatch")
        if not set(finding["source_rule_ids"]).issubset(set(disposition.get("source_rule_ids", []))):
            transfer_errors.append(f"{fid}:rule-id-mismatch")
        if disposition.get("included_in_return"):
            final_text = normalize(str(disposition.get("final_text") or ""))
            if finding["finding_type"] == "A":
                expected_numbers = set(finding["item_numbers"])
                disposition_numbers = {int(n) for n in disposition.get("item_numbers", []) if str(n).isdigit()}
                if not expected_numbers or not expected_numbers.issubset(disposition_numbers):
                    transfer_errors.append(f"{fid}:item-mismatch")
                for number in expected_numbers:
                    if not final_text or final_text not in normalize(draft_comments.get(number, "")):
                        transfer_errors.append(f"{fid}:missing-from-04-item-{number}")
            else:
                document_name = normalize(finding["document"] or "")
                if normalize(str(disposition.get("document") or "")) != document_name:
                    transfer_errors.append(f"{fid}:document-mismatch")
                if not final_text or final_text not in normalize(draft_b) or document_name not in normalize(draft_b):
                    transfer_errors.append(f"{fid}:missing-from-04-B")
        else:
            route = disposition.get("return_route")
            if route not in non_return_routes or fid not in factcheck:
                routing_errors.append(f"{fid}:route={route}:recorded_in_06={fid in factcheck}")
    add(checks, "transfer.findings_to_04", not transfer_errors, f"errors={transfer_errors}")
    add(checks, "transfer.non_return_recorded", not routing_errors, f"errors={routing_errors}")

    document = Document(args.word)
    add(checks, "word.table_count", len(document.tables) == 2, f"actual={len(document.tables)}")
    if len(document.tables) >= 2:
        dims = [(len(t.rows), len(t.columns)) for t in document.tables[:2]]
        add(checks, "word.table_shape", dims == [(52, 2), (1, 1)], f"actual={dims}")
        word_a = {row_item_number(row.cells[0].text): row.cells[1].text for row in document.tables[0].rows if row_item_number(row.cells[0].text)}
        a_mismatch = [number for number, comment in draft_comments.items() if normalize(word_a.get(number, "")) != normalize(comment)]
        add(checks, "transfer.04_to_word_a", not a_mismatch, f"mismatched_items={a_mismatch}")
        word_footer = document.tables[1].cell(0, 0).text
        add(checks, "transfer.04_to_word_b", normalize(word_footer) == normalize(draft_footer), f"exact_match={normalize(word_footer) == normalize(draft_footer)}")
        non_action_filled: list[int] = []
        for row in document.tables[0].rows:
            item_no = row_item_number(row.cells[0].text)
            if item_no not in draft_comments and len(row.cells) > 1 and normalize(row.cells[1].text):
                non_action_filled.append(item_no or -1)
        add(checks, "word.non_action_blank", not non_action_filled, f"filled={non_action_filled}")
        forbidden = ["読んだら消してください", "前回の申請では"]
        word_text = normalize("\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells))
        residue = [value for value in forbidden if value in word_text]
        add(checks, "word.no_standard_text_residue", not residue, f"residue={residue}")
    flags = docx_xml_flags(args.word)
    add(checks, "word.no_comments", flags["comment_parts"] == 0, str(flags))
    add(checks, "word.no_tracked_changes", flags["tracked_change_tags"] == 0, str(flags))
    expected_name = expected_word_name(draft, args.word.name)
    add(checks, "word.filename", expected_name is not None and args.word.name.casefold() == expected_name.casefold(), f"expected={expected_name}; actual={args.word.name}")
    return {"passed": all(c["passed"] for c in checks), "checks": checks}


def write_report(result: dict[str, Any], json_path: Path, md_path: Path) -> None:
    json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = ["# プレチェック機械検証結果", "", f"総合結果：{'PASS' if result['passed'] else 'FAIL'}", "", "| 検証 | 結果 | 詳細 |", "|---|---|---|"]
    for check in result["checks"]:
        detail = str(check["detail"]).replace("|", "\\|").replace("\n", " ")
        lines.append(f"| {check['id']} | {'PASS' if check['passed'] else 'FAIL'} | {detail} |")
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Deterministically verify precheck transfers and return DOCX structure.")
    parser.add_argument("--rules", type=Path, default=default_rules_dir())
    parser.add_argument("--timeline", type=Path, required=True, help="reviewed 00_timeline_facts.json")
    parser.add_argument("--ledger", type=Path, required=True)
    parser.add_argument("--review", type=Path, required=True, help="03 review markdown")
    parser.add_argument("--draft", type=Path, required=True, help="04 draft markdown")
    parser.add_argument("--factcheck", type=Path, required=True, help="06 fact-check markdown")
    parser.add_argument("--word", type=Path, required=True)
    parser.add_argument("--json-out", type=Path, required=True)
    parser.add_argument("--md-out", type=Path, required=True)
    args = parser.parse_args()
    result = verify(args)
    write_report(result, args.json_out, args.md_out)
    print(json.dumps({"passed": result["passed"], "failed": [c["id"] for c in result["checks"] if not c["passed"]]}, ensure_ascii=False))
    return 0 if result["passed"] else 1


if __name__ == "__main__":
    sys.exit(main())
