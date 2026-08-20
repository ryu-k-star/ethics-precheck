from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


EXTRACTOR_VERSION = "1.0.0"
SOURCE_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".zip"}
EXCLUDED_DIR_PATTERNS = (
    re.compile(r"^99_プレチェック出力(?:_|$)"),
    re.compile(r"^tmp$", re.I),
    re.compile(r"^output(?:s)?$", re.I),
    re.compile(r"^cache$", re.I),
    re.compile(r"^\.precheck-cache$", re.I),
    re.compile(r"^\.git$", re.I),
)
DATE_PATTERN = re.compile(
    r"(?P<year>20\d{2})\s*[年/.-]\s*(?P<month>1[0-2]|0?[1-9])\s*[月/.-]\s*(?P<day>3[01]|[12]\d|0?[1-9])\s*日?"
)
MONTH_PATTERN = re.compile(r"(?P<months>\d{1,3})\s*(?:か|ヶ|ケ|箇)?月")
TARGET_PATTERN = re.compile(r"(?P<count>\d{1,6})\s*(?:例|名|人)")
TIMELINE_TERMS = (
    "研究期間", "実施期間", "研究終了", "登録期間", "エントリー", "追跡", "観察期間",
    "解析", "統計解析", "症例数", "対象者数", "予定対象者数",
)
RULE_FILE_GROUPS = (
    ("R1_項目間整合ルール.md", "r1-cross-document-consistency.md"),
    ("R2_必須記載チェックリスト.md", "r2-required-content.md"),
)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def is_excluded(relative: Path) -> bool:
    return any(any(pattern.search(part) for pattern in EXCLUDED_DIR_PATTERNS) for part in relative.parts[:-1])


def discover_sources(study: Path) -> tuple[list[Path], list[str]]:
    sources: list[Path] = []
    warnings: list[str] = []
    for path in sorted(study.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SOURCE_SUFFIXES:
            continue
        rel = path.relative_to(study)
        if is_excluded(rel):
            if path.suffix.lower() in {".pdf", ".docx"}:
                warnings.append(f"excluded-source-like-file: {rel}")
            continue
        sources.append(path)
    return sources, warnings


def extract_pdf(path: Path) -> dict[str, Any]:
    reader = PdfReader(str(path))
    pages = []
    for number, page in enumerate(reader.pages, 1):
        pages.append({"page": number, "text": page.extract_text() or ""})
    return {"kind": "pdf", "pages": pages}


def extract_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    blocks: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        if paragraph.text.strip():
            blocks.append(f"[P{index}] {paragraph.text}")
    for table_no, table in enumerate(document.tables, 1):
        for row_no, row in enumerate(table.rows, 1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            blocks.append(f"[T{table_no}R{row_no}] " + " | ".join(cells))
    return {"kind": "docx", "pages": [{"page": None, "text": "\n".join(blocks)}]}


def extract_plain(path: Path) -> dict[str, Any]:
    return {"kind": path.suffix.lower().lstrip("."), "pages": [{"page": None, "text": path.read_text(encoding="utf-8", errors="replace")}]}


def extract_source(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".md", ".csv"}:
        return extract_plain(path)
    return {"kind": suffix.lstrip("."), "pages": [], "note": "binary inventory only"}


def cache_key(file_hash: str) -> str:
    return hashlib.sha256(f"{EXTRACTOR_VERSION}:{file_hash}".encode()).hexdigest()


def get_extraction(path: Path, file_hash: str, cache_dir: Path, force: bool) -> tuple[dict[str, Any], bool]:
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / f"{cache_key(file_hash)}.json"
    if cache_file.exists() and not force:
        return json.loads(cache_file.read_text(encoding="utf-8")), True
    payload = extract_source(path)
    payload.update({"extractor_version": EXTRACTOR_VERSION, "sha256": file_hash})
    cache_file.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload, False


def compact_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def timeline_evidence(extractions: list[dict[str, Any]]) -> dict[str, Any]:
    evidence: list[dict[str, Any]] = []
    for item in extractions:
        for page in item["extraction"].get("pages", []):
            lines = (page.get("text") or "").splitlines()
            for line_no, raw in enumerate(lines, 1):
                line = compact_line(raw)
                if not line or not any(term in line for term in TIMELINE_TERMS):
                    continue
                start = max(0, line_no - 2)
                end = min(len(lines), line_no + 2)
                context = compact_line(" ".join(lines[start:end]))
                dates = [m.group(0) for m in DATE_PATTERN.finditer(context)]
                months = [int(m.group("months")) for m in MONTH_PATTERN.finditer(context)]
                targets = [int(m.group("count")) for m in TARGET_PATTERN.finditer(context)]
                evidence.append({
                    "source": item["path"], "page": page.get("page"), "line": line_no,
                    "text": line[:500], "context": context[:1000], "dates": dates, "months": months, "targets": targets,
                })
    return {
        "schema_version": 1,
        "status": "requires_reviewer_confirmation",
        "mandatory_questions": [
            "研究開始日と研究終了日は全資料で一致するか",
            "明記された研究期間（月数）は年月日の差と一致するか",
            "最終登録後の追跡期間を研究終了日までに確保できるか",
            "追跡終了後から統計解析終了までの期間を確保できるか",
            "目標症例数と登録見込みから登録完了時期は現実的か",
        ],
        "evidence": evidence,
        "review_fields": {
            "research_start": None, "research_end": None, "stated_total_months": None,
            "target_count": None, "expected_enrollment_end": None, "followup_months": None,
            "analysis_months": None, "latest_followup_end": None, "latest_analysis_end": None,
            "analysis_margin_days": None, "finding": None, "assumptions": [],
        },
    }


def rule_ids(rules_dir: Path) -> list[str]:
    found: list[str] = []
    pattern = re.compile(r"^\|\s*((?:G|C)\d+-[0-9A-Za-z]+)\s*\|")
    for alternatives in RULE_FILE_GROUPS:
        path = next((rules_dir / name for name in alternatives if (rules_dir / name).is_file()), None)
        if path is None:
            raise FileNotFoundError(f"Missing rule file under {rules_dir}: one of {alternatives}")
        for line in path.read_text(encoding="utf-8").splitlines():
            match = pattern.match(line)
            if match and match.group(1) not in found:
                found.append(match.group(1))
    return found


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract one study once and create a compact precheck fact pack.")
    parser.add_argument("study", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--cache", type=Path, default=Path(".precheck-cache"))
    parser.add_argument("--rules", type=Path, default=Path("rules"))
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args()
    study = args.study.resolve()
    out = args.out.resolve()
    out.mkdir(parents=True, exist_ok=True)
    sources, warnings = discover_sources(study)
    if not sources:
        raise SystemExit(f"No source files found under {study}")

    manifest_files: list[dict[str, Any]] = []
    extracted: list[dict[str, Any]] = []
    for path in sources:
        stat = path.stat()
        file_hash = sha256_file(path)
        extraction, hit = get_extraction(path, file_hash, args.cache.resolve(), args.force)
        rel = str(path.relative_to(study))
        manifest_files.append({
            "path": rel, "suffix": path.suffix.lower(), "size": stat.st_size,
            "mtime_ns": stat.st_mtime_ns, "sha256": file_hash, "cache_hit": hit,
        })
        extracted.append({"path": rel, "extraction": extraction})

    manifest = {
        "schema_version": 1, "created_at": utc_now(), "study": str(study),
        "extractor_version": EXTRACTOR_VERSION, "excluded_directory_policy": [p.pattern for p in EXCLUDED_DIR_PATTERNS],
        "warnings": warnings, "files": manifest_files,
    }
    (out / "00_source_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    (out / "00_extracted_pages.json").write_text(json.dumps(extracted, ensure_ascii=False, indent=2), encoding="utf-8")
    timeline = timeline_evidence(extracted)
    (out / "00_timeline_facts.json").write_text(json.dumps(timeline, ensure_ascii=False, indent=2), encoding="utf-8")
    ledger = {
        "schema_version": 2, "study": study.name,
        "allowed_statuses": ["na", "satisfied", "candidate", "human"],
        "allowed_evidence_statuses": ["confirmed", "inferred", "human"],
        "allowed_return_routes": [
            "applicant_return", "attachment_return", "researcher_profile",
            "secretariat_internal", "human_decision", "resolved_elsewhere",
        ],
        "rules": [
            {
                "rule_id": rid, "status": None, "evidence": [], "note": "",
                "root_issue_id": None, "linked_items": [],
            }
            for rid in rule_ids(args.rules.resolve())
        ],
        "dispositions": [],
        "disposition_schema": {
            "finding_id": "案件内で一意のID",
            "source_rule_ids": ["G1-1"],
            "item_numbers": [1],
            "evidence_status": "confirmed|inferred|human",
            "return_route": "applicant_return|attachment_return|researcher_profile|secretariat_internal|human_decision|resolved_elsewhere",
            "action_owner": "申請者・各研究者本人・事務局・最終確認者など",
            "included_in_return": True,
            "disposition_reason": "採用・集約・別経路対応・保留等の理由",
            "root_issue_id": "同じ原因から派生する指摘を束ねるID",
        },
    }
    (out / "00_rule_ledger.json").write_text(json.dumps(ledger, ensure_ascii=False, indent=2), encoding="utf-8")
    summary = {
        "study": study.name, "source_count": len(sources), "cache_hits": sum(bool(f["cache_hit"]) for f in manifest_files),
        "timeline_evidence_count": len(timeline["evidence"]), "rule_count": len(ledger["rules"]), "warnings": warnings,
    }
    (out / "00_fact_pack.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
