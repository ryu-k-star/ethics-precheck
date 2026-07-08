from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def split_bottom_sections(text: str) -> tuple[str, str]:
    marker = "A．倫理審査申請システム質疑事項"
    idx = text.find(marker)
    if idx == -1:
        return text, ""
    return text[:idx].strip(), text[idx:].strip()


def extract_title(text: str) -> str:
    for line in text.splitlines():
        if line.startswith("研究課題名"):
            return line.strip()
    return "研究課題名【】"


def parse_markdown_table(text: str) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    in_table = False
    for line in text.splitlines():
        if line.startswith("| 項目 | 対応要否 |"):
            in_table = True
            continue
        if not in_table:
            continue
        if line.startswith("|---"):
            continue
        if not line.startswith("|"):
            if rows:
                break
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if len(cells) < 3:
            continue
        rows.append((cells[0], cells[1], cells[2].replace("<br>", "\n")))
    return rows


def clear_cell(cell) -> None:
    cell.text = ""


def build_docx(template: Path, draft: Path, output: Path) -> None:
    shutil.copy2(template, output)
    doc = Document(output)
    doc.styles["Normal"].font.name = "Yu Gothic"
    doc.styles["Normal"].font.size = Pt(10.5)

    draft_text = read_text(draft)
    top_text, bottom_text = split_bottom_sections(draft_text)
    title = extract_title(top_text)

    if len(doc.paragraphs) > 1:
        doc.paragraphs[1].text = title

    rows = parse_markdown_table(draft_text)
    if not rows:
        raise ValueError("No A-section markdown table found in draft.")
    if len(doc.tables) < 2:
        raise ValueError("Template must contain at least two tables.")

    table = doc.tables[0]
    for row_idx in range(len(table.rows)):
        clear_cell(table.cell(row_idx, 1))

    max_rows = min(len(rows), len(table.rows))
    for idx in range(max_rows):
        _item, status, comment = rows[idx]
        if status == "要対応" and comment:
            table.cell(idx, 1).text = comment

    doc.tables[1].cell(0, 0).text = bottom_text
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Create a return DOCX while preserving the official template tables."
    )
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--draft", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build_docx(args.template, args.draft, args.output)


if __name__ == "__main__":
    main()
