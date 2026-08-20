from __future__ import annotations

import argparse
import importlib.util
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "codex-skill"


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


class PrecheckV2PipelineTests(unittest.TestCase):
    def test_scan_reuses_hash_cache_and_excludes_prior_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            study = base / "study-001"
            study.mkdir()
            (study / "plan.txt").write_text(
                "1. 研究課題名 テスト研究\n研究期間 2026年8月1日から2027年3月31日まで\n"
                "予定対象者数 38名\n統計解析期間は未記載\n作成日 20XX年X月X日\n"
                "［試料・情報］を選択してください\n",
                encoding="utf-8",
            )
            roster = Document()
            table = roster.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "所属"
            table.cell(0, 1).text = "氏名"
            roster.save(study / "researchers.docx")
            excluded = study / "99_プレチェック出力_old"
            excluded.mkdir()
            Document().save(excluded / "old.docx")
            output = study / "99_プレチェック出力_v2"
            cache = base / "cache"
            command = [
                sys.executable,
                str(SKILL / "scripts" / "precheck_scan.py"),
                str(study),
                "--out",
                str(output),
                "--cache",
                str(cache),
            ]

            first = subprocess.run(command, capture_output=True, text=True)
            self.assertEqual(first.returncode, 0, first.stderr)
            second = subprocess.run(command, capture_output=True, text=True)
            self.assertEqual(second.returncode, 0, second.stderr)
            first_summary = json.loads(first.stdout)
            second_summary = json.loads(second.stdout)
            ledger = json.loads((output / "00_rule_ledger.json").read_text(encoding="utf-8"))
            manifest = json.loads((output / "00_source_manifest.json").read_text(encoding="utf-8"))
            r3 = json.loads((output / "00_r3_hits.json").read_text(encoding="utf-8"))
            fact_pack = json.loads((output / "00_fact_pack.json").read_text(encoding="utf-8"))

            self.assertEqual(first_summary["source_count"], 2)
            self.assertEqual(first_summary["cache_hits"], 0)
            self.assertEqual(second_summary["cache_hits"], 2)
            self.assertNotIn("fact_candidates", first_summary)
            self.assertNotIn("［試料", first.stdout)
            self.assertEqual(len(ledger["rules"]), 108)
            self.assertTrue(any("old.docx" in warning for warning in manifest["warnings"]))
            self.assertTrue(any(hit["rule_id"].startswith("T1-") for hit in r3["hits"]))
            self.assertTrue(any(hit["rule_id"] == "T1-4" for hit in r3["hits"]))
            self.assertTrue(any(hit["rule_id"] == "T1-13" for hit in r3["hits"]))
            self.assertIn("study_titles", fact_pack["fact_candidates"])
            self.assertIn(1, fact_pack["application_items"]["detected"])
            self.assertTrue(any(hit["rule_id"] == "T1-4" for hit in fact_pack["r3_hits"]))

    def test_verifier_rejects_dropped_swapped_and_missing_findings(self) -> None:
        verifier = load_module("precheck_verify", SKILL / "scripts" / "precheck_verify.py")
        self.assertEqual(verifier.default_rules_dir().resolve(), (SKILL / "references" / "rules").resolve())
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            ledger = {
                "allowed_evidence_statuses": ["confirmed", "inferred", "human"],
                "allowed_return_routes": [
                    "applicant_return",
                    "attachment_return",
                    "researcher_profile",
                    "secretariat_internal",
                    "human_decision",
                    "resolved_elsewhere",
                ],
                "rules": [
                    {"rule_id": rule_id, "status": "satisfied"}
                    for rule_id in verifier.expected_rule_ids(SKILL / "references" / "rules")
                ],
                "dispositions": [
                    {
                        "finding_id": "A-001", "finding_type": "A", "source_rule_ids": ["G6-1"],
                        "item_numbers": [17], "document": None, "evidence_status": "confirmed",
                        "return_route": "applicant_return", "action_owner": "申請者", "included_in_return": True,
                        "final_text": "登録期間を明記してください。", "disposition_reason": "申請者修正", "root_issue_id": "ROOT-1",
                    },
                    {
                        "finding_id": "A-002", "finding_type": "A", "source_rule_ids": ["G6-2", "G6-7"],
                        "item_numbers": [17], "document": None, "evidence_status": "confirmed",
                        "return_route": "applicant_return", "action_owner": "申請者", "included_in_return": True,
                        "final_text": "追跡・解析期間を整理してください。", "disposition_reason": "申請者修正", "root_issue_id": "ROOT-1",
                    },
                    {
                        "finding_id": "B-001", "finding_type": "B", "source_rule_ids": ["C1-1"],
                        "item_numbers": [], "document": "説明文書", "evidence_status": "confirmed",
                        "return_route": "attachment_return", "action_owner": "申請者", "included_in_return": True,
                        "final_text": "説明文書を修正して再アップロードしてください。", "disposition_reason": "添付修正", "root_issue_id": "ROOT-2",
                    },
                ],
            }
            ledger_path = base / "ledger.json"
            ledger_path.write_text(json.dumps(ledger), encoding="utf-8")
            review = base / "03.md"
            review.write_text(
                "## A欄候補\n\n"
                "| Finding ID | ルールID | 指摘先項目 | 根拠頁 | 内容 |\n"
                "|---|---|---|---|---|\n"
                "| A-001 | G6-1 | 17 | 4 | 登録期間を確認 |\n"
                "| A-002 | G6-2,G6-7 | 17 | 4,9 | 追跡・解析期間を確認 |\n\n"
                "## B欄候補\n\n"
                "| Finding ID | ルールID | 資料 | 根拠頁 | 内容 |\n"
                "|---|---|---|---|---|\n"
                "| B-001 | C1-1 | 説明文書 | 9 | 説明文書の期間を修正 |\n",
                encoding="utf-8",
            )
            draft = base / "04.md"
            rows = ["研究課題名【テスト研究】", "", "| 項目 | 対応要否 | 指摘内容 |", "|---|---|---|"]
            for item in range(1, 52):
                status = "要対応" if item == 17 else "対応不要"
                comment = "登録期間を明記してください。追跡・解析期間を整理してください。" if item == 17 else ""
                rows.append(f"| {item}. item | {status} | {comment} |")
            rows.append("| 連絡担当者 | 対応不要 |  |")
            rows.extend(["", "A．倫理審査申請システム質疑事項", "B．修正が必要な添付書類の再アップロード", "説明文書を修正して再アップロードしてください。", "C．軽微な文言修正等を行った添付書類の再アップロード", "なし"])
            draft.write_text("\n".join(rows), encoding="utf-8")
            factcheck = base / "06.md"
            factcheck.write_text("確認済み", encoding="utf-8")
            timeline = base / "timeline.json"
            timeline.write_text(json.dumps({"status": "reviewed", "review_fields": {
                "research_start": "2026-08-01", "research_end": "2027-03-31", "stated_total_months": 8,
                "target_count": 38, "expected_enrollment_end": None, "followup_months": None,
                "analysis_months": None, "latest_followup_end": None, "latest_analysis_end": None,
                "analysis_margin_days": None, "finding": "期間を要整理", "assumptions": [],
            }}), encoding="utf-8")

            document = Document()
            main = document.add_table(rows=52, cols=2)
            for index, row in enumerate(main.rows):
                label = f"{index + 1}. item" if index < 51 else "連絡担当者"
                row.cells[0].text = label
                if index == 16:
                    row.cells[1].text = "登録期間を明記してください。追跡・解析期間を整理してください。"
            footer = document.add_table(rows=1, cols=1)
            footer.cell(0, 0).text = "A．倫理審査申請システム質疑事項\nB．修正が必要な添付書類の再アップロード\n説明文書を修正して再アップロードしてください。\nC．軽微な文言修正等を行った添付書類の再アップロード\nなし"
            word = base / "★倫理申請_修正対応依頼_テスト研究_0820.docx"
            document.save(word)

            result = verifier.verify(
                argparse.Namespace(
                    rules=SKILL / "references" / "rules",
                    timeline=timeline,
                    ledger=ledger_path,
                    review=review,
                    draft=draft,
                    factcheck=factcheck,
                    word=word,
                )
            )
            self.assertTrue(result["passed"])

            original_draft = draft.read_text(encoding="utf-8")
            draft.write_text(original_draft.replace("追跡・解析期間を整理してください。", ""), encoding="utf-8")
            dropped = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=review, draft=draft, factcheck=factcheck, word=word))
            self.assertFalse(next(c for c in dropped["checks"] if c["id"] == "transfer.findings_to_04")["passed"])
            draft.write_text(original_draft, encoding="utf-8")

            swapped_doc = Document(word)
            swapped_doc.tables[0].rows[0].cells[1].text = swapped_doc.tables[0].rows[16].cells[1].text
            swapped_doc.tables[0].rows[16].cells[1].text = ""
            swapped_doc.save(word)
            swapped = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=review, draft=draft, factcheck=factcheck, word=word))
            self.assertFalse(next(c for c in swapped["checks"] if c["id"] == "transfer.04_to_word_a")["passed"])

            document.save(word)
            no_b = original_draft.replace("説明文書を修正して再アップロードしてください。", "")
            draft.write_text(no_b, encoding="utf-8")
            missing_b = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=review, draft=draft, factcheck=factcheck, word=word))
            self.assertFalse(next(c for c in missing_b["checks"] if c["id"] == "transfer.findings_to_04")["passed"])

            moved_to_c = original_draft.replace(
                "説明文書を修正して再アップロードしてください。\nC．",
                "C．",
            ).replace("C．軽微な文言修正等を行った添付書類の再アップロード\nなし", "C．軽微な文言修正等を行った添付書類の再アップロード\n説明文書を修正して再アップロードしてください。")
            draft.write_text(moved_to_c, encoding="utf-8")
            moved_doc = Document(word)
            moved_doc.tables[1].cell(0, 0).text = "\n".join(moved_to_c.splitlines()[-4:])
            moved_doc.save(word)
            moved = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=review, draft=draft, factcheck=factcheck, word=word))
            self.assertFalse(next(c for c in moved["checks"] if c["id"] == "transfer.findings_to_04")["passed"])

            draft.write_text(original_draft, encoding="utf-8")
            document.save(word)
            unknown_review = base / "03-unknown.md"
            unknown_review.write_text(review.read_text(encoding="utf-8").replace("G6-1 |", "G6-99 |", 1), encoding="utf-8")
            unknown = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=unknown_review, draft=draft, factcheck=factcheck, word=word))
            self.assertFalse(next(c for c in unknown["checks"] if c["id"] == "findings.schema")["passed"])

            wrong_name = base / "★倫理申請_修正対応依頼_別研究_0820.docx"
            document.save(wrong_name)
            wrong_filename = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=review, draft=draft, factcheck=factcheck, word=wrong_name))
            self.assertFalse(next(c for c in wrong_filename["checks"] if c["id"] == "word.filename")["passed"])

            draft.write_text(original_draft, encoding="utf-8")
            for row in ledger["rules"]:
                if row["rule_id"].startswith("G6-"):
                    row["status"] = "na"
            ledger_path.write_text(json.dumps(ledger), encoding="utf-8")
            invalid_g6 = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=review, draft=draft, factcheck=factcheck, word=word))
            self.assertFalse(next(c for c in invalid_g6["checks"] if c["id"] == "duration.core_applicable")["passed"])

            for row in ledger["rules"]:
                row["status"] = "satisfied"
            ledger["dispositions"] = []
            ledger_path.write_text(json.dumps(ledger), encoding="utf-8")
            empty_review = base / "03-empty.md"
            empty_review.write_text("## A欄候補\n\n| Finding ID | ルールID | 指摘先項目 | 根拠頁 | 内容 |\n|---|---|---|---|---|\n\n## B欄候補\n\n| Finding ID | ルールID | 資料 | 根拠頁 | 内容 |\n|---|---|---|---|---|\n", encoding="utf-8")
            empty_rows = ["研究課題名【指摘なし】", "", "| 項目 | 対応要否 | 指摘内容 |", "|---|---|---|"]
            empty_rows.extend(f"| {item}. item | 対応不要 |  |" for item in range(1, 52))
            empty_rows.extend(["| 連絡担当者 | 対応不要 |  |", "", "A．倫理審査申請システム質疑事項", "B．修正が必要な添付書類の再アップロード", "なし", "C．軽微な文言修正等を行った添付書類の再アップロード", "なし"])
            empty_draft = base / "04-empty.md"
            empty_draft.write_text("\n".join(empty_rows), encoding="utf-8")
            empty_doc = Document()
            empty_main = empty_doc.add_table(rows=52, cols=2)
            for index, row in enumerate(empty_main.rows):
                row.cells[0].text = f"{index + 1}. item" if index < 51 else "連絡担当者"
            empty_footer = empty_doc.add_table(rows=1, cols=1)
            empty_footer.cell(0, 0).text = "A．倫理審査申請システム質疑事項\nB．修正が必要な添付書類の再アップロード\nなし\nC．軽微な文言修正等を行った添付書類の再アップロード\nなし"
            empty_word = base / "★倫理申請_修正対応依頼_指摘なし_0820.docx"
            empty_doc.save(empty_word)
            empty_result = verifier.verify(argparse.Namespace(rules=SKILL / "references" / "rules", timeline=timeline, ledger=ledger_path, review=empty_review, draft=empty_draft, factcheck=factcheck, word=empty_word))
            self.assertTrue(empty_result["passed"], empty_result)


if __name__ == "__main__":
    unittest.main()
