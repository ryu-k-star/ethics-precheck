from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document


def safe_name(path: Path) -> str:
    name = re.sub(r'[<>:"/\\|?*\s]+', "_", path.stem).strip("_")
    return name[:120] or "document"


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    parts.append(f"# {path.name}")
    parts.append("")
    parts.append("## Paragraphs")
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        if text:
            parts.append(f"[P{i}] {text}")
    table_count = 0
    for table in doc.tables:
        table_count += 1
        parts.append("")
        parts.append(f"## Table {table_count}")
        for r_idx, row in enumerate(table.rows, 1):
            cells = [" ".join(c.text.split()) for c in row.cells]
            parts.append(f"[T{table_count}R{r_idx}] " + " | ".join(cells))
    return "\n".join(parts).strip() + "\n"


def extract_pdf(path: Path) -> str:
    parts: list[str] = [f"# {path.name}", ""]
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_no = page.page_number
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            parts.append(f"## Page {page_no}")
            parts.append(text.strip())
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for t_idx, table in enumerate(tables, 1):
                parts.append(f"### Page {page_no} Table {t_idx}")
                for row in table:
                    cells = [" ".join((c or "").split()) for c in row]
                    parts.append(" | ".join(cells))
            parts.append("")
    return "\n".join(parts).strip() + "\n"


def project_dirs(root: Path) -> list[Path]:
    dirs: list[Path] = []
    for child in root.iterdir():
        if child.is_dir():
            nested = [p for p in child.iterdir() if p.is_dir()]
            if len(nested) == 1 and not any(child.glob("*.docx")) and not any(child.glob("*.pdf")):
                dirs.append(nested[0])
            else:
                dirs.append(child)
    return sorted(dirs)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("root", type=Path)
    parser.add_argument("--out", type=Path, default=Path("tmp/extracted"))
    args = parser.parse_args()

    args.out.mkdir(parents=True, exist_ok=True)
    manifest = []
    for project in project_dirs(args.root):
        project_out = args.out / project.parent.name / project.name
        project_out.mkdir(parents=True, exist_ok=True)
        files = sorted([p for p in project.rglob("*") if p.is_file()])
        listing = []
        for path in files:
            rel = path.relative_to(project)
            listing.append(
                {
                    "path": str(rel),
                    "size": path.stat().st_size,
                    "suffix": path.suffix.lower(),
                }
            )
            try:
                if path.suffix.lower() == ".docx":
                    text = extract_docx(path)
                    out_file = project_out / f"{safe_name(path)}.md"
                    out_file.write_text(text, encoding="utf-8")
                elif path.suffix.lower() == ".pdf":
                    text = extract_pdf(path)
                    out_file = project_out / f"{safe_name(path)}.md"
                    out_file.write_text(text, encoding="utf-8")
                elif path.suffix.lower() == ".zip":
                    with zipfile.ZipFile(path) as zf:
                        names = zf.namelist()
                    (project_out / f"{safe_name(path)}_zip_contents.txt").write_text(
                        "\n".join(names), encoding="utf-8"
                    )
            except Exception as exc:
                (project_out / f"{safe_name(path)}.error.txt").write_text(
                    f"{type(exc).__name__}: {exc}\n", encoding="utf-8"
                )
        (project_out / "00_file_listing.json").write_text(
            json.dumps(listing, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        manifest.append({"project": str(project), "output": str(project_out), "files": listing})

    (args.out / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main()
